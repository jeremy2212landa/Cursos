from docx import Document
from docx.shared import Cm

# Creación del documento
document = Document()

# Añadimos un titulo al documento, a nivel 0
document.add_heading('Documento creado con Python', 0)

# Añadimos un párrafo
p = document.add_paragraph('El contenido de los párrafos se añadir en varias líneas. ')
p.add_run('Pudiéndose configurar que el texto tenga formato tipo ')
p.add_run('negrita').bold = True
p.add_run(' o ')
p.add_run('itálica.').italic = True

# Para indicar subtitulo se indica el nivel 1
document.add_heading('Subtitulo', level=1)

document.add_paragraph('Ahora se puede crear una enumeración')
document.add_paragraph('Uno', style='List Number')
document.add_paragraph('Dos', style='List Number')
document.add_paragraph('Tres', style='List Number')

document.add_paragraph('O viñetas')
document.add_paragraph('Manzana', style='List Bullet') 
document.add_paragraph('Pera', style='List Bullet')
document.add_paragraph('Naranja', style='List Bullet')

# Imágenes
document.add_heading('Imágenes', level=1)
document.add_picture('/home/kattaleya/Documentos/develop/py/Python_course/scripts/eff.jpg', width=Cm(5))

# Tablas
document.add_heading('Tablas', level=1)

data = (('Manzana', 12), ('Pera', 5), ('Naranja', 12))

table = document.add_table(rows=1, cols=2)

table.rows[0].cells[0].text = 'Fruta'
table.rows[0].cells[1].text = 'Cantidad'

for prod, numbr in data:
    row_cells = table.add_row().cells
    row_cells[0].text = prod
    row_cells[1].text = str(numbr)

document.save('ejemplo.docx')